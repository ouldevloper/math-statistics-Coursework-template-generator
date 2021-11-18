from asyncio.tasks import wait
from math import log, pi
from os import path, remove
import requests as req
import json as js
import asyncio
from docx import Document
import pprint
import math
import matplotlib.pyplot as plt
import io
import random as rand
Xdata = [168, 158, 173, 165, 164, 161, 150, 155, 153, 182,
         165, 184, 169, 170, 185, 165, 155, 164, 174, 145,
         166, 181, 151, 158, 170, 184, 153, 158, 151, 164,
         170, 168, 151, 163, 156, 170, 184, 164, 164, 164,
         147, 171, 165, 165, 156, 164, 181, 181, 181, 168,
         155, 165, 157, 170, 157, 184, 170, 183, 171, 184,
         184, 150, 189, 155, 147, 163, 183, 158, 185, 178,
         145, 148, 163, 158, 157, 162, 164, 174, 150, 164,
         164, 155, 156, 182, 150, 178, 168, 156, 163, 163,
         164, 167, 149, 165, 151, 155, 153, 162, 155, 151,
         148, 170, 156, 155, 184, 162, 185, 170, 183, 164,
         168, 147, 164, 170, 185, 181, 156, 156, 169, 168,
         151, 170, 158, 168, 148, 172, 184, 149, 170, 171,
         147, 167, 162, 153, 182, 169, 150, 163, 153, 164,
         160, 160, 150, 173, 160, 174, 161, 150, 183, 168,
         181, 173, 156, 181, 170, 170, 157, 181, 181, 153,
         176, 164, 156, 165, 156, 184, 170, 166, 182, 156,
         149, 181, 170, 170, 169, 150, 176, 158, 185, 145,
         161, 184, 165, 164, 157, 166, 158, 176, 165, 183,
         150, 171, 178, 185, 157, 182, 171, 164, 161, 181]

Ydata = [85, 91, 80, 89, 86, 88, 93, 92, 83, 88, 97, 87, 87,
         86, 95, 93, 96, 93, 94, 87, 90, 99, 97, 86, 96, 85, 83,
         83, 93, 97, 92, 81, 98, 97, 93, 85, 99, 87, 87, 94, 80,
         89, 83, 97, 86, 88, 97, 86, 80, 92, 99, 88, 82, 92, 93,
         98, 85, 80, 84, 80, 95, 88, 82, 92, 82, 99, 82, 84, 95,
         91, 94, 85, 88, 91, 86, 94, 84, 98, 81, 80, 86, 88, 98,
         98, 81, 85, 82, 87, 87, 87, 82, 99, 96, 84, 99, 92, 87,
         86, 85, 91, 97, 83, 87, 94, 93, 92, 82, 86, 80, 87, 99,
         97, 92, 85, 84, 84, 85, 85, 82, 92, 80, 83, 80, 93, 93,
         84, 92, 99, 88, 93, 93, 98, 81, 95, 82, 91, 96, 88, 88,
         88, 86, 94, 97, 98, 94, 90, 93, 97, 83, 90, 92, 91, 80,
         96, 97, 90, 81, 85, 93, 97, 87, 86, 88, 87, 89, 81, 91,
         94, 88, 80, 99, 81, 98, 80, 94, 91, 96, 94, 93, 92, 91,
         88, 89, 94, 84, 84, 88, 97, 96, 90, 98, 84, 90, 83, 92,
         98, 80, 98, 92, 99]

_dataXY = []
filename = "test.docx"
table1 = []
table2 = []
table3 = []
table4 = []
table5 = []
table6 = []
table7 = []
h = 0
doc = Document()
xbbar = 0
db = 0
betab = 0
S = 0
Spow2 = 0


async def getresponse():
    file = open("file.txt", "rb+")
    data = file.read()
    await asyncio.sleep(1)
    file.close()
    if len(data) <= 0:
        try:
            data = req.get("https://api.covid19api.com/summary").content
            await asyncio.sleep(1)
            remove("data.txt") if path.exists("data.txt") else None
            file = open("data.txt", "wb+")
            file.write(data)
            file.close()
        except Exception as err:
            print("Error => " + str(err))
    return data.decode("utf8")


async def tojson(data):
    l = await asyncio.create_task(getresponse())
    d = js.loads(l)
    await asyncio.sleep(1)
    return d


async def sortlist(lst):
    data = sorted(lst, key=lambda x: x[0])
    asyncio.sleep(1)
    return data


async def filldata():
    dd = await getresponse()
    json = await asyncio.create_task(tojson(dd))
    for line in json["Countries"]:
        # _dataZ.append(line['Contry'])
        _dataXY.append([line['TotalConfirmed'], line['TotalDeaths']])


def showData(lst):
    index_x = 0
    index_y = 0
    i = 0
    for dt in range(len(lst) // 10 if len(lst) % 10 == 0 else (len(lst) // 10)+1):
        last = 10 if index_x+10 < len(lst) else len(lst) % 10
        dd = lst[index_x:index_x+last]
        for ii in range(2):
            print("\nX " if ii == 0 else "\nY ", end="| ")
            for item in dd:
                print("%9d" % item[ii], end='  ')
        index_x += 10
        print("\n\n")


###########################################################################################

def writestep(num):
    p = doc.add_paragraph("")
    p.add_run("Step "+str(num)).bold = True
    doc.save(filename)


def generateData():
    for i in range(200):
        # num1 = rand.randint(145, 195)
        # num2 = rand.randint(80, 100)
        table1.append([Xdata[i], Ydata[i]])


def step1_2(lst):
    global tablenmuber
    index_x = 0
    index_y = 0
    i = 0
    for dt in range(len(lst) // 10 if len(lst) % 10 == 0 else (len(lst) // 10)+1):
        last = 10 if index_x+10 <= len(lst) else len(lst) % 10
        dd = lst[index_x:index_x+last]

        t = doc.add_table(rows=2, cols=11)

        t.rows[0].cells[0].text = "X"
        t.rows[1].cells[0].text = "Y"
        for index, item in enumerate(dd):
            t.rows[0].cells[index+1].text = str(item[0])
            t.rows[1].cells[index+1].text = str(item[1])
        # for ii in range(2):
        #    print("\nX " if ii == 0 else "\nY ", end="| ")
        #    for item in dd:
        #        print("%9d" % item[ii], end='  ')
        index_x += 10
        # print("")
        doc.add_paragraph("")
    doc.save(filename)


def step2(lst):
    table2 = sorted(lst, key=lambda x: x[0])


def step3():
    i = 0
    tmp = list(set(Xdata))
    table = 0
    t = None
    for index, item in enumerate(tmp):
        if index % 10 == 0:
            doc.add_paragraph("")
            i = 0
            t = doc.add_table(rows=4, cols=11)
            t.rows[0].cells[0].text = "i"
            t.rows[1].cells[0].text = "xi"
            t.rows[2].cells[0].text = "pi"
            t.rows[3].cells[0].text = "ni"
        table3.append([index+1, item, Xdata.count(item),
                      (Xdata.count(item)/len(table1))])
        t.rows[0].cells[i+1].text = str(index+1)
        t.rows[1].cells[i+1].text = str(item)
        t.rows[2].cells[i+1].text = str(Xdata.count(item))
        t.rows[3].cells[i+1].text = "%d/200" % Xdata.count(item)
        i += 1
    doc.save(filename)


def step4():
    global h
    h_ = (max(Xdata)-min(Xdata))/(1+3.28*log(200))
    h = math.ceil(h_)
    h_txt = "h = xmax-xmin/1+3.28*ln N  = %d-%d/1+3.28*ln %d = %f = %d" % (
        max(Xdata), min(Xdata), 200, h_, h)
    p = doc.add_paragraph("")
    p.add_run(h_txt).bold = True
    doc.save(filename)


def step5_hep_1(start, end):
    count = 0
    for item in set([item[0] for item in table1]):
        if item in range(start if start == min(Xdata) else start+1, end+1):
            count += Xdata.count(item)
    return count


def step5():
    global Xdata
    _data = set(Xdata)
    l = (max(_data)-min(_data))
    lst = l//h if l % h == 0 else (l//h)+1
    item = min(_data)
    doc.add_paragraph("")
    t = doc.add_table(rows=lst+1, cols=4)
    t.rows[0].cells[0].text = "i"
    t.rows[0].cells[1].text = "xi<X<=xi+1"
    t.rows[0].cells[2].text = "ni"
    t.rows[0].cells[3].text = "pi=ni/N"

    sumni = 0
    for index in range(lst):
        t.rows[index+1].cells[0].text = str(index+1)
        t.rows[index+1].cells[1].text = "%d-%d" % (item, item+h)
        tmp = step5_hep_1(item, item+h)
        item += h
        sumni += tmp
        t.rows[index+1].cells[2].text = "%d" % (tmp)
        t.rows[index+1].cells[3].text = "%d/200" % tmp
        table4.append([index+1, (item+item+h)/2, tmp, tmp/len(table1)])
    p = doc.add_paragraph("")
    p.add_run("sum ni = %d/200 = %d" % (sumni, (sumni//len(table1))))
    doc.save(filename)


X = []


def step6():
    doc.add_paragraph("")
    t = doc.add_table(rows=len(table4)+1, cols=2)
    t.rows[0].cells[0].text = "i"
    t.rows[0].cells[1].text = "F*(x)"
    lastnum = 0
    for index, item in enumerate([item[2] for item in table4]):
        if index == 0:
            t.rows[index+1].cells[0].text = "1"
            t.rows[index +
                   1].cells[1].text = "%d/%d" % (table4[0][2], len(table1))
            table5.append([1, item/len(table1)])
        else:
            t.rows[index+1].cells[0].text = str(index+1)
            t.rows[index+1].cells[1].text = "%d/%d + %d/%d = %d/%d" % (
                lastnum, len(Xdata), table4[index][2], len(Xdata), lastnum+table4[index][2], len(Xdata))
            table5.append([index+1, (lastnum+item)/len(table1)])
        lastnum += item
    doc.save('test.docx')


def step7():
    l = (max(Xdata)-min(Xdata))
    lst = l//h if l % h == 0 else (l//h)+1
    t = doc.add_table(rows=lst+1, cols=4)
    t.rows[0].cells[0].text = "i"
    t.rows[0].cells[1].text = "Xi"
    t.rows[0].cells[2].text = "pi"
    t.rows[0].cells[3].text = "ni / h*N"
    for i, item in enumerate(range(min(Xdata), max(Xdata)+1, math.ceil(h))):
        t.rows[i+1].cells[0].text = "%d" % (i+1)
        t.rows[i+1].cells[1].text = "%f" % ((item+item+3)/2)
        t.rows[i+1].cells[2].text = "%f" % (table4[i][3]/len(Xdata))
        t.rows[i+1].cells[3].text = "%f" % (table4[i][2]/(h*len(Xdata)))
        table6.append([i+1, (item+item+3)/2, table4[i][3] /
                      len(Xdata), table4[i][2]/(h*len(Xdata))])
    doc.save(filename)


def step8():
    # Y = list(range(min(Xdata), max(Xdata), len(Y)//len(Fx)))
    Y = [item[1] for item in table5]
    X = [item[1] for item in table4]
    plt.plot(X, Y, "-o")
    plt.xlabel("X")
    plt.ylabel("F*(x)")
    plt.savefig("step8.png")
    doc.add_picture("step8.png")
    doc.save(filename)
    plt.close()


def step9():
    global table5
    global table6
    Y = [item[1] for item in table5]
    X = [item[2] for item in table6]
    plt.plot(Y, X, "-o")
    plt.savefig("step9.png")
    p = doc.add_paragraph("")
    p.add_run("Step 9").bold = True
    doc.add_picture("step9.png")
    doc.save("test.docx")
    plt.close()


def step10():
    global xbbar
    global db
    global betab
    sm = sum([item[1]*item[2] for item in table4])
    xbbar = sm/len(table1)
    db = sum([(item[2]*(item[1]-xbbar)**2)/len(table1) for item in table4])
    betab = math.sqrt(db)
    doc.add_paragraph("")
    doc.add_paragraph("xb(bar) = %.3f/%d = %f" %
                      (sm, len(table1), xbbar))
    doc.add_paragraph("")
    doc.add_paragraph("db = ni*(xi-xbbar)^2/N  = %f " % db)
    doc.add_paragraph("")
    doc.add_paragraph("betab = sqrt(db) = sqrt(%f) = %f" % (db, betab))
    doc.add_paragraph("")
    X = [item[1] for item in table4]
    Y = [(item/(len(table1)*h)) for item in [item[2] for item in table4]]
    plt.plot(X, Y)
    plt.savefig("step10.png")
    doc.add_picture("step10.png")

    S = math.sqrt(len(table1) / (len(table1)-1))*betab
    Spow2 = (len(table1) / (len(table1)-1))*db

    doc.add_paragraph("")
    doc.add_paragraph("S = sqrt(N/N-1)*betab = %f " % S)
    doc.add_paragraph("")

    doc.add_paragraph("")
    doc.add_paragraph("S^2 = (N/N-1)*db = %f " % Spow2)
    doc.add_paragraph("")

    doc.save(filename)
    plt.close()


def step11():
    t = doc.add_table(rows=len(table6)+1, cols=7)
    t.rows[0].cells[0].text = "xi"
    t.rows[0].cells[1].text = "xi-xb"
    t.rows[0].cells[2].text = "ui = (xi-xb)/betab"
    t.rows[0].cells[3].text = "laplaca(ui)"
    t.rows[0].cells[4].text = "ni' = (nh/betab)*laplaca"
    t.rows[0].cells[5].text = "ni'"
    t.rows[0].cells[6].text = "pi*'"
    sum = 0
    for index, item in enumerate(table6):
        col1 = item[1]-xbbar
        col2 = col1/betab
        col3 = (math.exp(1) **
                (-(((item[1]-xbbar)/betab)**2)/2))*(1/math.sqrt(2*math.pi))
        col4 = (len(table1)*h/betab)*col3

        # if math.ceil(col3)-col4 > 0.4 else math.floor(col4)
        col5 = math.ceil(col4)
        col6 = col5/len(table1)
        table7.append([item[1], col1, col2, col3, col4, col5, col6])
        t.rows[index+1].cells[0].text = str(item[1])
        t.rows[index+1].cells[1].text = f"{col1:.3f}"
        t.rows[index+1].cells[2].text = f"{col2:.3f}"
        t.rows[index+1].cells[3].text = f"{col3:.3f}"
        t.rows[index + 1].cells[4].text = f"{col4:.3f}"
        t.rows[index+1].cells[5].text = f"{col5}"
        t.rows[index+1].cells[6].text = f"{col6:.3f}"
        sum += col5
    p = doc.add_paragraph("")
    p.add_run("Sum n'i = %d" % sum)
    doc.save("test.docx")


def step12():
    global table5
    global table6
    plt.clf()
    plt.close()
    X = [item[1] for item in table6]
    Y = [item[3] for item in table4]
    Z = [item[6] for item in table7]
    plt.plot(X, Y, "-s", label="l1")
    plt.plot(X, Z, "-s", label="l2")
    plt.legend()
    plt.savefig("step12.png")
    p = doc.add_paragraph("")
    p.add_run("Step 12").bold = True
    doc.add_picture("step12.png")
    doc.save("test.docx")
    plt.close()


#########################################
generateData()
writestep(1)
step1_2(table1)
writestep(2)
step2(table1)
step1_2(table2)
writestep(3)
step3()
writestep(4)
step4()
writestep(5)
step5()
writestep(6)
step6()
writestep(7)
step7()
writestep(8)
step8()
writestep(9)
step9()
writestep(10)
step10()
writestep(11)
step11()
writestep(12)
step12()


async def main():
    await asyncio.create_task(getresponse())
    await asyncio.create_task(filldata())
    # showData(_dataXY)
    data = await asyncio.create_task(sortlist(_dataXY))
    showData(data)


# asyncio.run(main())

# lst = [[9, 4], [5, 3], [2, 5], [0, 1], [8, 0]]
# srt = sorted(lst, key=lambda x: x[0])
# pprint.pprint(srt)
