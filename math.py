from asyncio.tasks import wait
from math import hypot, log, pi, sin
from os import path, remove
import requests as req
import json as js
import asyncio
from docx import Document
import pprint
from pprint import pprint
from scipy.integrate import quad
import math
import matplotlib.pyplot as plt
from docx.shared import Cm, Inches
import sys
import signal
import io
import random as rand


class CollectDataFromAPI:
    def __init__(self):
        pass


class MathStatistics:
    def __init__(self, filename="output.docx"):
        self.Xdata = [148, 155, 156, 156, 157, 157, 158, 159, 159, 159, 160, 160,
                      160, 161, 161, 161, 161, 161, 161, 162, 162, 162, 162, 162,
                      162, 163, 163, 163, 163, 163, 163, 163, 163, 163, 164,
                      164, 164, 164, 164, 164, 164, 164, 164, 164, 164, 164, 164,
                      164, 164, 165, 165, 165, 165, 165, 165, 165, 165, 165, 165,
                      165, 165, 165, 165, 165, 166, 166, 166, 166, 166, 166, 166,
                      166, 167, 167, 167, 167, 167, 167, 167, 167, 167, 167, 167,
                      167, 167, 167, 168, 168, 168, 168, 168, 168, 168, 168, 168,
                      168, 169, 169, 169, 169, 169, 169, 169, 169, 169, 169, 169,
                      169, 169, 169, 169, 170, 170, 170, 170, 170, 170, 170, 170,
                      170, 170, 170, 170, 170, 170, 170, 171, 171, 171, 171, 171,
                      171, 171, 171, 171, 171, 171, 172, 172, 172, 172, 172, 172,
                      172, 172, 172, 172, 172, 172, 173, 173, 173, 173, 173, 173,
                      173, 173, 173, 174, 174, 174, 174, 174, 174, 175, 175, 175,
                      175, 175, 175, 176, 176, 176, 176, 176, 176, 176, 176, 176,
                      176, 177, 178, 178, 178, 178, 179, 179, 179, 180, 180, 180,
                      181, 181, 182, 183, 184, 185, 186, 187, 190]

        self.Ydata = [87, 86, 85, 88, 82, 90, 91, 81, 83, 86, 83, 85,
                      87, 79, 82, 84, 84, 88, 91, 80, 89, 90, 92, 94,
                      100, 80, 88, 88, 88, 89, 91, 91, 92, 93, 83,
                      84, 84, 84, 85, 86, 87, 89, 89, 89, 90, 90, 91,
                      91, 97, 84, 85, 87, 87, 87, 87, 88, 90, 91, 91,
                      93, 94, 94, 94, 94, 82, 84, 84, 84, 87, 88, 89,
                      89, 81, 85, 85, 85, 86, 88, 88, 89, 89, 89, 90,
                      91, 91, 92, 81, 82, 83, 87, 88, 88, 90, 92, 93,
                      93, 79, 83, 87, 87, 87, 88, 88, 89, 91, 91, 91,
                      91, 91, 92, 92, 81, 86, 88, 88, 90, 90, 90, 90,
                      91, 91, 91, 92, 93, 95, 96, 83, 85, 89, 89, 90,
                      91, 91, 91, 92, 94, 97, 82, 87, 88, 89, 90, 91,
                      91, 94, 96, 98, 99, 99, 84, 87, 89, 89, 90, 90,
                      90, 91, 96, 86, 88, 90, 91, 96, 97, 85, 89, 90,
                      90, 91, 95, 82, 86, 87, 88, 90, 92, 93, 93, 93,
                      95, 87, 89, 90, 90, 91, 85, 85, 99, 85, 90, 98,
                      89, 92, 90, 90, 98, 91, 92, 86, 105]

        # self.Xdata = [168, 158, 173, 165, 164, 161, 150, 155, 153, 182,
        #              165, 184, 169, 170, 185, 165, 155, 164, 174, 145,
        #              166, 181, 151, 158, 170, 184, 153, 158, 151, 164,
        #              170, 168, 151, 163, 156, 170, 184, 164, 164, 164,
        #              147, 171, 165, 165, 156, 164, 181, 181, 181, 168,
        #              155, 165, 157, 170, 157, 184, 170, 183, 171, 184,
        #              184, 150, 189, 155, 147, 163, 183, 158, 185, 178,
        #              145, 148, 163, 158, 157, 162, 164, 174, 150, 164,
        #              164, 155, 156, 182, 150, 178, 168, 156, 163, 163,
        #              164, 167, 149, 165, 151, 155, 153, 162, 155, 151,
        #              148, 170, 156, 155, 184, 162, 185, 170, 183, 164,
        #              168, 147, 164, 170, 185, 181, 156, 156, 169, 168,
        #              151, 170, 158, 168, 148, 172, 184, 149, 170, 171,
        #              147, 167, 162, 153, 182, 169, 150, 163, 153, 164,
        #              160, 160, 150, 173, 160, 174, 161, 150, 183, 168,
        #              181, 173, 156, 181, 170, 170, 157, 181, 181, 153,
        #              176, 164, 156, 165, 156, 184, 170, 166, 182, 156,
        #              149, 181, 170, 170, 169, 150, 176, 158, 185, 145,
        #              161, 184, 165, 164, 157, 166, 158, 176, 165, 183,
        #              150, 171, 178, 185, 157, 182, 171, 164, 161, 181]
        # self.Ydata = [85, 91, 80, 89, 86, 88, 93, 92, 83, 88, 97, 87, 87,
        #              86, 95, 93, 96, 93, 94, 87, 90, 99, 97, 86, 96, 85, 83,
        #              83, 93, 97, 92, 81, 98, 97, 93, 85, 99, 87, 87, 94, 80,
        #              89, 83, 97, 86, 88, 97, 86, 80, 92, 99, 88, 82, 92, 93,
        #              98, 85, 80, 84, 80, 95, 88, 82, 92, 82, 99, 82, 84, 95,
        #              91, 94, 85, 88, 91, 86, 94, 84, 98, 81, 80, 86, 88, 98,
        #              98, 81, 85, 82, 87, 87, 87, 82, 99, 96, 84, 99, 92, 87,
        #              86, 85, 91, 97, 83, 87, 94, 93, 92, 82, 86, 80, 87, 99,
        #              97, 92, 85, 84, 84, 85, 85, 82, 92, 80, 83, 80, 93, 93,
        #              84, 92, 99, 88, 93, 93, 98, 81, 95, 82, 91, 96, 88, 88,
        #              88, 86, 94, 97, 98, 94, 90, 93, 97, 83, 90, 92, 91, 80,
        #              96, 97, 90, 81, 85, 93, 97, 87, 86, 88, 87, 89, 81, 91,
        #              94, 88, 80, 99, 81, 98, 80, 94, 91, 96, 94, 93, 92, 91,
        #              88, 89, 94, 84, 84, 88, 97, 96, 90, 98, 84, 90, 83, 92,
        #              98, 80, 98, 92, 99]
        self.laplac = {
            0.0: [0, 399, 789, 1197, 1595, 1994, 2392, 2790, 3188, 3586],
            0.1: [3983, 4380, 4776, 5172, 5567, 5962, 6356, 6749, 7142, 7535],
            0.2: [7926, 8317, 8706, 9095, 9483, 9871, 10257, 10642, 11026, 11409],
            0.3: [11791, 12172, 12552, 12930, 13307, 13683, 14058, 14431, 14803, 15173],
            0.4: [15542, 15910, 16276, 16640, 17003, 17364, 17724, 18082, 18439, 18793],
            0.5: [19146, 19497, 19847, 20194, 20540, 20884, 21226, 21566, 21904, 22240],
            0.6: [22575, 22907, 23237, 23565, 23891, 24215, 24537, 24857, 25175, 25490],
            0.7: [25804, 26115, 26424, 26730, 27035, 27337, 27637, 27935, 28230, 28524],
            0.8: [28814, 29103, 29389, 29673, 29955, 30234, 30511, 30785, 31057, 31327],
            0.9: [31594, 31859, 32121, 32381, 32639, 32894, 33147, 33398, 33646, 33891],
            1.0: [34134, 34375, 34614, 34850, 35083, 35314, 35543, 35769, 35993, 36214],
            1.1: [36433, 36650, 36864, 37076, 37286, 37493, 37698, 38000, 38100, 38298],
            1.2: [38493, 38686, 38877, 39065, 39251, 39435, 39617, 39796, 39973, 40147],
            1.3: [40320, 40490, 40658, 40824, 40988, 41149, 41308, 41466, 41621, 41774],
            1.4: [41924, 42073, 42220, 42364, 42507, 42647, 42786, 42922, 43056, 43189],
            1.5: [43319, 43448, 43574, 43699, 43822, 43943, 44062, 44179, 44295, 44408],
            1.6: [44520, 44630, 44738, 44845, 44950, 45053, 45154, 45254, 45352, 45449],
            1.7: [45543, 45637, 45728, 45818, 45907, 45994, 46080, 46164, 46246, 46327],
            1.8: [46407, 46485, 46562, 46638, 46712, 46784, 46856, 46926, 46995, 47062],
            1.9: [47128, 47193, 47257, 47320, 47381, 47441, 47500, 47558, 47615, 47670],
            2.0: [47725, 47778, 47831, 47882, 47932, 47982, 48030, 48077, 48124, 48169],
            2.1: [48214, 48257, 48300, 48341, 48382, 48422, 48461, 48500, 48537, 48574],
            2.2: [48610, 48645, 48679, 48713, 48745, 48778, 48806, 48840, 48870, 48899],
            2.3: [48928, 48956, 48983, 49010, 49036, 49061, 49086, 49111, 49134, 49158],
            2.4: [49180, 49202, 49224, 49245, 49266, 49286, 49305, 49324, 49343, 49361],
            2.5: [49379, 49396, 49413, 49430, 49446, 49461, 49477, 49492, 49506, 49520],
            2.6: [49534, 49547, 49560, 49573, 49585, 49598, 49609, 49621, 49632, 49643],
            2.7: [49653, 49664, 49674, 49683, 49693, 49702, 49711, 49720, 49728, 49736],
            2.8: [49744, 49752, 49760, 49767, 49774, 49781, 49788, 49795, 49801, 49807],
            2.9: [49813, 49819, 49825, 49831, 49836, 49841, 49846, 49851, 49856, 49861],
            3.0: [49865, 49869, 49874, 49878, 49882, 49886, 49889, 49893, 49896, 49900],
            3.1: [49903, 49906, 49910, 49913, 49916, 49918, 49921, 49924, 49926, 49929],
            3.2: [49931, 49934, 49936, 49938, 49940, 49942, 49944, 49946, 49948, 49950, 0],
            3.3: [49952, 49953, 49955, 49957, 49958, 49960, 49961, 49962, 49964, 49965, 0],
            3.4: [49966, 49968, 49969, 49970, 49971, 49972, 49973, 49974, 49975, 49976, 0],
            3.5: [49977, 49978, 49978, 49979, 49980, 49981, 49981, 49982, 49983, 49983, 0],
            3.6: [49984, 49985, 49985, 49986, 49986, 49987, 49987, 49988, 49988, 49989, 0],
            3.7: [49989, 49990, 49990, 49990, 49991, 49991, 49992, 49992, 49992, 49992, 0],
            3.8: [49993, 49993, 49993, 49994, 49994, 49994, 49994, 49995, 49995, 49995, 0],
            3.9: [49995, 49995, 49996, 49996, 49996, 49996, 49996, 49996, 49997, 49997, 0],
            4.0: [49997, 0, 0, 0, 0, 0, 0, 0, 0, 0]}
        self.filename = filename
        self.table1 = []
        self.table2 = []
        self.table3 = []
        self.table4 = []
        self.table5 = []
        self.table6 = []
        self.table7 = []
        self.table8 = []
        self.table9 = []
        self.table10 = []
        self.h = 0
        self.N = len(self.Xdata)
        self.doc = Document()
        self.xbbar = 0
        self.db = 0
        self.segmab = 0
        self.S = 0
        self.Spow2 = 0
        self.hy = 0
        # self.xbar = 0
        # self.ybar = 0
        # self.xbar2 = 0
        # self.ybar2 = 0

    def showData(self, lst):
        index_x = 0
        index_y = 0
        i = 0
        for dt in range(len(lst) // 10 if len(lst) % 10 == 0 else (len(lst) // 10)+1):
            last = 10 if index_x+10 < len(lst) else len(lst) % 10
            dd = lst[index_x:index_x+last]
            for ii in range(2):
                print("\nX " if ii == 0 else "\nY ", end="| ")
                for item in dd:
                    print("%9d" % item[ii], end='  ')
            index_x += 10
            print("\n\n")

###########################################################################################

    def writestep(self, num):
        print(f"Step {num}")
        p = self.doc.add_paragraph("")
        p.add_run("Step "+str(num)).bold = True
        self.doc.save(self.filename)

    def generateData(self):
        for i in range(len(self.Xdata)):
            # num1 = rand.randint(145, 195)
            # num2 = rand.randint(80, 100)
            self.table1.append([self.Xdata[i], self.Ydata[i]])

    def step1(self):
        index_x = 0
        index_y = 0
        i = 0
        for dt in range(self.N // 10 if self.N % 10 == 0 else (self.N // 10)+1):
            last = 10 if index_x + \
                10 <= self.N else self.N % 10
            dd = self.table1[index_x:index_x+last]
            t = self.doc.add_table(rows=2, cols=11)
            t.rows[0].cells[0].text = "X"
            t.rows[1].cells[0].text = "Y"
            for index, item in enumerate(dd):
                t.rows[0].cells[index+1].text = str(item[0])
                t.rows[1].cells[index+1].text = str(item[1])
            # for ii in range(2):
            #    print("\nX " if ii == 0 else "\nY ", end="| ")
            #    for item in dd:
            #        print("%9d" % item[ii], end='  ')
            index_x += 10
            # print("")
            self.doc.add_paragraph("")
        self.doc.save(self.filename)

    def step2(self):
        self.table2 = sorted(self.table1, key=lambda x: x[0])
        index_x = 0
        index_y = 0
        i = 0
        for dt in range(len(self.table2) // 10 if len(self.table2) % 10 == 0 else (len(self.table2) // 10)+1):
            last = 10 if index_x + \
                10 <= len(self.table2) else len(self.table2) % 10
            dd = self.table2[index_x:index_x+last]
            t = self.doc.add_table(rows=2, cols=11)
            t.rows[0].cells[0].text = "X"
            t.rows[1].cells[0].text = "Y"
            for index, item in enumerate(dd):
                t.rows[0].cells[index+1].text = str(item[0])
                t.rows[1].cells[index+1].text = str(item[1])
            # for ii in range(2):
            #    print("\nX " if ii == 0 else "\nY ", end="| ")
            #    for item in dd:
            #        print("%9d" % item[ii], end='  ')
            index_x += 10
            # print("")
            self.doc.add_paragraph("")
        self.doc.save(self.filename)

    def step3(self):
        i = 0
        tmp = list(set(self.Xdata))
        table = 0
        t = None
        for index, item in enumerate(tmp):
            if index % 10 == 0:
                self.doc.add_paragraph("")
                i = 0
                t = self.doc.add_table(rows=4, cols=11)
                t.rows[0].cells[0].text = "i"
                t.rows[1].cells[0].text = "xi"
                t.rows[2].cells[0].text = "pi"
                t.rows[3].cells[0].text = "ni"
            t.rows[0].cells[i+1].text = str(index+1)
            t.rows[1].cells[i+1].text = str(item)
            t.rows[2].cells[i+1].text = str(self.Xdata.count(item))
            t.rows[3].cells[i +
                            1].text = "%d/%d" % (self.Xdata.count(item), self.N)
            self.table3.append([index+1, item, self.Xdata.count(item),
                                (self.Xdata.count(item)/self.N)])
            i += 1
        self.doc.save(self.filename)

    def step4(self):
        h_ = (max(self.Xdata)-min(self.Xdata))/(1+3.28*log(len(self.Xdata)))
        self.h = math.ceil(h_)
        h_txt = "h = xmax-xmin/1+3.28*ln N  = %d-%d/1+3.28*ln %d = %f = %d" % (
            max(self.Xdata), min(self.Xdata), self.N, h_, self.h)
        p = self.doc.add_paragraph("")
        p.add_run(h_txt).bold = True
        self.doc.save(self.filename)

    def step5_hep_1(self, start, end):
        count = 0
        for item in set([item[0] for item in self.table1]):
            if item in range(start if start == min(self.Xdata) else start+1, end+1):
                count += self.Xdata.count(item)
        return count

    def step5(self):
        _data = set(self.Xdata)
        l = (max(_data)-min(_data))
        lst = l//self.h if l % self.h == 0 else (l//self.h)+1
        item = min(_data)
        self.doc.add_paragraph("")
        t = self.doc.add_table(rows=lst+1, cols=4)
        t.rows[0].cells[0].text = "i"
        t.rows[0].cells[1].text = "xi<X<=xi+1"
        t.rows[0].cells[2].text = "ni"
        t.rows[0].cells[3].text = "pi=ni/N"

        sumni = 0
        for index in range(lst):
            t.rows[index+1].cells[0].text = str(index+1)
            t.rows[index+1].cells[1].text = "%d-%d" % (item, item+self.h)
            tmp = self.step5_hep_1(item, item+self.h)
            item += self.h
            sumni += tmp
            t.rows[index+1].cells[2].text = "%d" % (tmp)
            t.rows[index+1].cells[3].text = "%d/%d" % (tmp, self.N)
            self.table4.append(
                [index+1, (item+item+self.h)/2, tmp, tmp/self.N])
        p = self.doc.add_paragraph("")
        p.add_run("sum ni = %d/%d = %d" %
                  (sumni, self.N, (sumni//self.N)))
        self.doc.save(self.filename)

    def step6(self):
        t = self.doc.add_table(rows=len(self.table4)+1, cols=2)
        t.rows[0].cells[0].text = "i"
        t.rows[0].cells[1].text = "F*(x)"
        lastnum = 0
        for index, item in enumerate([item[2] for item in self.table4]):
            if index == 0:
                t.rows[index+1].cells[0].text = "1"
                t.rows[index +
                       1].cells[1].text = "%d/%d" % (self.table4[0][2], self.N)
                self.table5.append([1, item/self.N])
            else:
                t.rows[index+1].cells[0].text = str(index+1)
                t.rows[index+1].cells[1].text = "%d/%d + %d/%d = %d/%d" % (
                    lastnum, len(self.Xdata), self.table4[index][2], len(self.Xdata), lastnum+self.table4[index][2], len(self.Xdata))
                self.table5.append(
                    [index+1, (lastnum+item)/self.N])
            lastnum += item
        self.doc.save(self.filename)

    def step7(self):
        l = (max(self.Xdata)-min(self.Xdata))
        lst = l//self.h if l % self.h == 0 else (l//self.h)+1
        t = self.doc.add_table(rows=lst+1, cols=4)
        t.rows[0].cells[0].text = "i"
        t.rows[0].cells[1].text = "Xi"
        t.rows[0].cells[2].text = "pi"
        t.rows[0].cells[3].text = "ni / h*N"
        # pprint.pprint(self.table4)
        res = min(self.Xdata)
        for i, item in enumerate([item[2] for item in self.table4]):
            t.rows[i+1].cells[0].text = "%d" % (i+1)
            t.rows[i+1].cells[1].text = f"{(res*2+self.h)/2:.1f}"
            t.rows[i+1].cells[2].text = f"{item/self.N:.3f}"
            t.rows[i+1].cells[3].text = f"{item/(self.h*self.N):.3f}"
            self.table6.append([i+1, (res*2+self.h)/2, item /
                               self.N, self.table4[i][2]/(self.h*self.N)])
            res += self.h

        self.doc.save(self.filename)

    def step8(self):
        # Y = list(range(min(Xdata), max(Xdata), len(Y)//len(Fx)))
        Y = [item[1] for item in self.table5]
        X = [item[1] for item in self.table4]
        plt.plot(X, Y, "-o")
        plt.xlabel("X")
        plt.ylabel("F*(x)")
        plt.savefig("step8.png")
        self.doc.add_picture("step8.png")
        self.doc.save(self.filename)
        plt.close()

    def step9(self):
        Y = [item[1] for item in self.table6]
        X = [item[2] for item in self.table6]
        plt.plot(Y, X, "-o")
        plt.savefig("step9.png")
        p = self.doc.add_paragraph("")
        p.add_run("Step 9").bold = True
        self.doc.add_picture("step9.png")
        self.doc.save(self.filename)
        plt.close()

    def step10(self):
        sm = sum([item[1]*item[2] for item in self.table3])
        self.xbbar = sm/self.N
        self.db = sum([(item[2]*(item[1]**2)) /
                      self.N for item in self.table3])-self.xbbar**2
        self.segmab = math.sqrt(self.db)
        self.doc.add_paragraph("")
        self.doc.add_paragraph("xb(bar) = %.3f/%d = %f" %
                               (sm, self.N, self.xbbar))
        self.doc.add_paragraph("")
        self.doc.add_paragraph("db = ni*(xi-xbbar)^2/N  = %f " % self.db)
        self.doc.add_paragraph("")
        self.doc.add_paragraph(
            "segmab = sqrt(db) = sqrt(%f) = %f" % (self.db, self.segmab))
        self.doc.add_paragraph("")
        X = [item[1] for item in self.table4]
        Y = [(item/(self.N*self.h))
             for item in [item[2] for item in self.table4]]
        plt.plot(X, Y)
        plt.savefig("step10.png")
        self.doc.add_picture("step10.png")

        self.S = math.sqrt(self.N / (self.N-1))*self.segmab
        self.Spow2 = (self.N / (self.N-1))*self.db

        self.doc.add_paragraph("")
        self.doc.add_paragraph("S = sqrt(N/N-1)*segmab = %f " % self.S)
        self.doc.add_paragraph("")

        self.doc.add_paragraph("")
        self.doc.add_paragraph("S^2 = (N/N-1)*db = %f " % self.Spow2)
        self.doc.add_paragraph("")

        self.doc.save(self.filename)
        # print("S:", self.S)
        # print("S^2:", self.Spow2)
        # print("N:", self.N)
        # print("db:", self.db)
        # print("xbar:", self.xbbar)
        # print("segmab:", self.segmab)
        plt.close()

    def step11(self):
        def floatcalc(val: float) -> float:
            fmt = f"{val}".split('.')
            if(int(fmt[1][0]) >= 5):
                return math.ceil(val)
            else:
                return int(val)
        t = self.doc.add_table(rows=len(self.table6)+1, cols=7)
        t.rows[0].cells[0].text = "xi"
        t.rows[0].cells[1].text = "xi-xb"
        t.rows[0].cells[2].text = "ui = (xi-xb)/segmab"
        t.rows[0].cells[3].text = "laplaca(ui)"
        t.rows[0].cells[4].text = "ni' = (nh/segmab)*laplaca"
        t.rows[0].cells[5].text = "ni'"
        t.rows[0].cells[6].text = "pi*'"
        sum = 0
        for index, item in enumerate([item[1] for item in self.table6]):
            col1 = item-self.xbbar
            col2 = col1/self.segmab
            col3 = (math.exp(1) ** (-(((item-self.xbbar)/self.segmab)**2)/2)
                    )*(1/math.sqrt(2*math.pi))
            col4 = (self.N*self.h/self.segmab)*col3

            # if math.ceil(col3)-col4 > 0.4 else math.floor(col4)
            col5 = floatcalc(col4)
            col6 = col5/self.N
            self.table7.append(
                [item, col1, col2, col3, col4, col5, col6])
            t.rows[index+1].cells[0].text = str(item)
            t.rows[index+1].cells[1].text = f"{col1:.3f}"
            t.rows[index+1].cells[2].text = f"{col2:.3f}"
            t.rows[index+1].cells[3].text = f"{col3:.3f}"
            t.rows[index + 1].cells[4].text = f"{col4:.3f}"
            t.rows[index+1].cells[5].text = f"{col5}"
            t.rows[index+1].cells[6].text = f"{col6:.3f}"
            sum += col5
        p = self.doc.add_paragraph("")
        p.add_run("Sum n'i = %d" % sum)
        self.doc.save(self.filename)

    def step12(self):
        xi = ['-∞']+list(map(str, [item[0] for item in self.table7]))+['+∞']
        t = self.doc.add_table(rows=len(xi)+1, cols=9)
        t.rows[0].cells[0].text = "i"
        t.rows[0].cells[1].text = "xi - xi+1"
        t.rows[0].cells[2].text = "ni"
        t.rows[0].cells[3].text = "Ф(zi)"
        t.rows[0].cells[4].text = "Fn(xi)"
        t.rows[0].cells[5].text = "Fn(xi+1)"
        t.rows[0].cells[6].text = "pmi"
        t.rows[0].cells[7].text = "nim = pmin"
        t.rows[0].cells[8].text = "(ni - nim)^2/nm i"

        print(f"S = {self.S}")
        print(f"Xbbar = {self.xbbar}")

        # def calczi(i):
        #    return ((self.table6[i-1][1]-self.xbbar)/self.S if i > 0 else 0) if i < len(self.table6) else 1

        def calclaplac(i):
            zi = ((self.table6[i-1][1]-self.xbbar)/self.S if i >
                  0 else 0) if i < len(self.table6) else 1
            laplac = ""
            if i <= 0:
                laplac = "-0.5"
            else:
                sign = "-" if f"{zi}".count('-') == 1 else ""
                _zi = f"{zi:.10f}".replace('-', '')
                tst = _zi.split('.')
                laplac = self.laplac[float(
                    f"{tst[0]}.{tst[1][0]}")][int(tst[1][1])]
                laplac = f"{sign}0.{laplac}"
            return float(f"{float(laplac):.4f}")
        _ni_ = 0
        _nmi_ = 0
        for index in range(len(xi)-1):
            col1 = index
            col2 = "%s÷%s" % (xi[index], xi[index+1])
            col3 = self.table4[index-1][2] if index != 0 else 0
            col4 = calclaplac(index)
            col5 = calclaplac(index)+0.5
            col6 = calclaplac(index+1)+0.5
            col7 = float((col6)-(col5))
            col8 = col7*self.N
            col9 = 0
            _nmi_ += col8
            _ni_ += col3
            if _nmi_ > 5:
                col9 = (_ni_-_nmi_)**2/_nmi_
                t.rows[index+1].cells[8].text = str(col9)
                # print(_ni_, "  =>   ", _nmi_, " ==>", (_ni_-_nmi_)**2/_nmi_)
                _nmi_ = 0
                _ni_ = 0
            else:
                col9 = 0
                t.rows[index+1].cells[8].text = "-"
            print(f"{col1}\t{col2}\t{col3}\t{col4}\t{col5}\t{col6}\t{col7}\t{col8}")
            # t.rows[index+1].cells[0].text = str(col1)
            # t.rows[index+1].cells[1].text = str(col2)
            # t.rows[index+1].cells[2].text = str(col3)
            # t.rows[index+1].cells[3].text = f"{col4:.3f}"
            # t.rows[index+1].cells[4].text = f"{col5:.4f}"
            # t.rows[index+1].cells[5].text = f"{col6:.4f}"
            # t.rows[index+1].cells[6].text = f"{col7:.4f}"
            # t.rows[index+1].cells[7].text = f"{col8:.4f}"

        plt.clf()
        plt.close()
        X = [item[1] for item in self.table6]
        Y = [item[3] for item in self.table4]
        Z = [item[6] for item in self.table7]
        plt.plot(X, Y, "-s", label="l1")
        plt.plot(X, Z, "-s", label="l2")
        plt.legend()
        plt.savefig("step12.png")
        p = self.doc.add_paragraph("")
        p.add_run("Step 12").bold = True
        self.doc.add_picture("step12.png")
        self.doc.save(self.filename)
        plt.close()

    def get_count_for_step14(self, val, _val):
        tmp = list(filter(lambda x: x[0] == val, self.table1))
        return list(map(lambda x: x[1], tmp)).count(_val)

    def step13(self):
        lenydata = len(set(self.Ydata))
        lenxdata = len(set(self.Xdata))
        t = self.doc.add_table(rows=lenxdata+2, cols=lenydata+2)
        t.columns[0].width = Inches(0.5)

        for i in range(1, lenydata+2):
            t.columns[i].width = Inches(0.3)
        t.columns[lenydata + 1].width = Inches(0.5)
        t.rows[0].cells[0].text = "X/Y"

        for i, data in enumerate(set(self.Ydata)):
            t.rows[0].cells[i+1].text = f"{data}"
        t.rows[0].cells[lenydata+1].text = "nxi"

        for i, data in enumerate(set(self.Xdata)):
            t.rows[i+1].cells[0].text = f"{data}"
        t.rows[lenxdata+1].cells[0].text = "yii"

        tmp_x_table = [0]*lenydata
        for index, data in enumerate(set(self.Xdata)):
            tmp_table = []
            for _index, _data in enumerate(set(self.Ydata)):
                t.rows[index+1].cells[_index +
                                      1].text = f"{self.get_count_for_step14(data, _data)}"
                xx = self.get_count_for_step14(data, _data)
                tmp_table.append(xx)
                tmp_x_table[_index] += xx
            x = sum(tmp_table)
            t.rows[index+1].cells[lenydata +
                                  1].text = f"{x}"
            tmp_table.append(x)
            self.table9.append(tmp_table)
        self.table9.append(tmp_x_table)

        p = self.doc.add_paragraph("")
        p.add_run("xbar=1/n... ybar=....").bold = True
        p.add_run(f"xbar={(1/self.N)*sum(self.Xdata)}").bold = True
        p.add_run(f"ybar={(1/self.N)*sum(self.Ydata)}").bold = True
        p.add_run(
            f"x^2bar={(1/self.N)*sum([item**2 for item in self.Xdata])}").bold = True
        p.add_run(
            f"y^2bar={(1/self.N)*sum([item**2 for item in self.Ydata])}").bold = True
        p.add_run(
            f"xybar={(1/self.N)*sum([item*_item  for item,_item in zip(self.Xdata,self.Ydata)])}").bold = True

        plt.clf()
        plt.figure()
        plt.plot(self.Xdata, self.Ydata, 'ro')
        plt.savefig("step13.png")
        self.doc.add_picture("step13.png")

        for i, data in enumerate(tmp_x_table):
            t.rows[lenxdata+1].cells[i+1].text = f"{data}"
        t.rows[lenxdata+1].cells[lenydata +
                                 1].text = f"{sum(tmp_x_table)}"
        self.hy = min(list(filter(lambda x: x > 1, tmp_x_table)))
        self.doc.add_paragraph(f"we take Y randomly : h={self.hy}")
        Y = set(self.Ydata)
        X = [item[0] for item in self.table7]
        __Y = list(range(min(Y), max(Y), self.hy))
        t = self.doc.add_table(rows=len(__Y)+3, cols=len(X)+3)
        t.columns[1].width = Inches(0.5)
        t.rows[0].cells[0].text = "N"
        t.rows[1].cells[1].text = "Y/X"
        t.rows[1].cells[len(X)+2].text = "nyi"
        t.rows[len(__Y)+2].cells[1].text = "nxi"
        for i, item in enumerate(X):
            t.rows[0].cells[i+2].text = f"{i+1}"
            t.rows[1].cells[i+2].text = f"{item}"

        for i, item in enumerate(__Y):
            t.rows[i+2].cells[0].text = f"{i+1}"
            t.rows[i+2].cells[1].text = f"{item}"
        # step15
        X_ = set([item[0] for item in self.table1])
        Y_ = set([item[1] for item in self.table1])

        def help(lst, startx, endx, starty, endy):
            _ll = list(filter(lambda x: x[0] in range(startx, endx, 1), lst))
            var = list(filter(lambda x: x[1] in range(starty, endy, 1), _ll))
            return len(var)
        for y, _ in enumerate(range(min(Y_), max(Y_), self.hy)):
            for i, _ in enumerate(range(min(X_), max(X_), self.h)):
                # print(help(self.table1, i, i+self.h if i+self.h < len(self.table1) else len(self.table1),
                #      y, y+self.hy if y+self.hy < len(self.table1[0]) else len(self.table1[0])), end=' ')
                firstx = i
                lastx = i+self.h  # if i + self.h < len(X_) else len(X_)
                firsty = y
                lasty = y+self.hy  # if y + self.hy < len(Y_) else len(Y_)
                data = help(self.table1, firstx, lastx, firsty, lasty)
                print(data, end=' ')
                # rows=len(__Y)+3, cols=len(X)+3
                t.rows[y+2].cells[i+2].text = f"{data}"
            print()
        self.doc.save(self.filename)


    #########################################
mt = MathStatistics()
mt.generateData()
mt.writestep(1)
mt.step1()
mt.writestep(2)
mt.step2()
mt.writestep(3)
mt.step3()
mt.writestep(4)
mt.step4()
mt.writestep(5)
mt.step5()
mt.writestep(6)
mt.step6()
mt.writestep(7)
mt.step7()
mt.writestep(8)
mt.step8()
mt.writestep(9)
mt.step9()
mt.writestep(10)
mt.step10()
mt.writestep(11)
mt.step11()
mt.writestep(12)
mt.step12()
mt.writestep(13)
mt.step13()

#
